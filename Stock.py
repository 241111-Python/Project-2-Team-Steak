from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import humanize

class Stock:
    def __init__(self, name):
        self.name = name
        self.data = []

    def addData(self, data):
        self.data.append(data)

    # Method returning the most recent data
    def latest_data(self):
        return self.data[-1]

    # Method retrieving the data from a specific date selected by the user
    def specific_date_data(self, date_string):
        date_object = datetime.strptime(date_string, "%Y-%m-%d")
        # Check if the day selected is a weekend
        if (date_object.weekday() >= 5):
            print("Please enter a valid weekday! Stock markets close on weekends.")
            return
        else:
            exists = False
            for i in self.data:
                # Look for chosen date in the available data
                if (i.date == date_string):
                    exists = True
                    print(i)
                    break
            if (exists == False):
                print("The data is not currently available for this date!\n")
    
    # Method retrieving the data for a specified range of dates
    def range_date_data(self, date_start_string, date_end_string, is_yearly = False):

        date_start_object = datetime.strptime(date_start_string, "%Y-%m-%d")
        date_end_object = datetime.strptime(date_end_string, "%Y-%m-%d")

        if (date_end_object < date_start_object):
            print("Please make sure your end date comes AFTER your start date!")
            return

        start_exists = False
        end_exists = False
        in_range = False
        # To calculate percentage change throughout the time period
        start_price = 0
        end_price = 0
        # For average closing and volume calculations
        total_closing = 0
        total_volume = 0
        days = 0
        # Sum of (close - open) over the time period
        close_open_diff = 0

        lowest_price = 9999999
        highest_price = 0

        # Lists for plotting
        dates = []
        closing_prices = []
        opening_prices = []

        for i in self.data:
            date_object = datetime.strptime(i.date, "%Y-%m-%d")
            # If we have a match with the start date or exceeded the start date
            if ((i.date == date_start_string) or ((date_object > date_start_object) and (start_exists == False))):
                start_price = i.close
                in_range = True
                start_exists = True
            # If we have a match with the end date or exceeded it
            if ((i.date == date_end_string) or (date_object > date_end_object)):
                end_price = i.close
                in_range = False
                end_exists = True
                break
            # When in range (between start and end dates) -> collect data for analysis
            if (in_range == True):
                if (i.low < lowest_price):
                    lowest_price = i.low
                if (i.high > highest_price):
                    highest_price = i.high
                total_closing += i.close
                total_volume += i.volume
                days += 1
                close_open_diff += (i.close - i.open)
                dates.append(i.date)
                closing_prices.append(i.close)
                opening_prices.append(i.open)
        # Check if start and end dates were reached
        if ((start_exists == False) or (end_exists == False)):
            print("Please enter valid start and end dates!")
            return

        # Calculate stats
        percentage_change = ((end_price - start_price)/start_price)*100
        avg_closing = total_closing/days
        avg_volume = total_volume/days

        # Make the stats readable
        rounded_values = [round(var, 2) for var in [lowest_price, highest_price, avg_closing, percentage_change, close_open_diff]]
        readable_volume = humanize.intword(avg_volume)

        # Display the calculated stats over the time period
        print("\nLowest price is: " + str(rounded_values[0]))
        print("Highest price is: " + str(rounded_values[1]))
        print("Average closing price is: " + str(rounded_values[2]))
        print("Change in price is: " + str(rounded_values[3]) + "%")
        print("Closing vs opening price is: " + str(rounded_values[4]))
        print("Average volume over this time period is: " + str(readable_volume))
        print("=========================================================")

        # If the data is not yearly -> plot both opening and closing prices
        if (is_yearly == False):

            fig, ax1 = plt.subplots()

            # Plotting closing price vs date
            ax1.plot(dates, closing_prices, color='b', label='Closing Price')
            ax1.set_xlabel('Date')
            ax1.set_ylabel('Closing Price', color='b')
            ax1.tick_params(axis='y', labelcolor='b')

            # Plotting opening price vs date (on the same graph)
            ax2 = ax1.twinx()
            ax2.plot(dates, opening_prices, color='r', label='Volume')
            ax2.set_ylabel('Opening Price', color='r')
            ax2.tick_params(axis='y', labelcolor='r')

            plt.xticks([dates[0], dates[-1]])
            # Increase or decrease in price (arrow)
            if (percentage_change > 0):
                plt.title("Daily " + self.name +  " Opening and Closing Prices( ↑" + str(round(percentage_change,2)) + "%)")
            else:
                plt.title("Daily " + self.name +  " Opening and Closing Prices( ↓" + str(round(percentage_change,2)) + "%)")
            plt.xticks(rotation=45)
            plt.tight_layout()

        # If the data is yearly (too much data) -> Only plot closing prices
        else:
            plt.figure(figsize=(10,5))
            plt.plot(dates, closing_prices, color='b', label='Closing Price')
            if (percentage_change > 0):
                plt.title("Daily " + self.name +  " Opening and Closing Prices( ↑" + str(round(percentage_change,2)) + "%)")
            else:
                plt.title("Daily " + self.name +  " Opening and Closing Prices( ↓" + str(round(percentage_change,2)) + "%)")
            plt.xlabel('Date')
            plt.ylabel('Closing Price')
            plt.xticks([dates[0], dates[-1]])
            plt.tight_layout()

        # Save the output to a word (docx) document
        plt.savefig("graph.png")
        doc = Document()
        doc.add_paragraph("Retrieving data from " + date_start_string + " to " + date_end_string +
                          "\nLowest price is: " + str(rounded_values[0]) + 
                          "\nHighest price is: " + str(rounded_values[1]) + 
                          "\nAverage closing price is: " + str(rounded_values[2]) + 
                          "\nChange in price is: " + str(rounded_values[3]) + 
                          "%\nClosing vs opening price is: " + str(rounded_values[4]) + 
                          "\nAverage volume over this time period is: " + str(readable_volume) + "\n")
        doc.add_picture('graph.png', width=Inches(6))
        doc.save('statistics.docx')
        plt.show()

    # Method that calls range_date_data passing the year and quarter start and end dates
    def quarterly_data(self, year_string, quarter):
        if (quarter == 1):
            date_start_string = year_string + "-01-01"
            date_end_string = year_string + "-03-31"
            self.range_date_data(date_start_string, date_end_string)

        elif (quarter == 2):
            date_start_string = year_string + "-04-01"
            date_end_string = year_string + "-06-30"
            self.range_date_data(date_start_string, date_end_string)

        elif (quarter == 3):
            date_start_string = year_string + "-07-01"
            date_end_string = year_string + "-09-30"
            self.range_date_data(date_start_string, date_end_string)

        elif (quarter == 4):
            date_start_string = year_string + "-10-01"
            date_end_string = year_string + "-12-31"
            self.range_date_data(date_start_string, date_end_string)
        
        else:
            print("Please enter a valid quarter!")

    # Method that calls range_date_data passing the year start and end dates
    def yearly_data(self, year_string):
        date_start_string = year_string + "-01-01"
        date_end_string = year_string + "-12-31"
        self.range_date_data(date_start_string, date_end_string, True)

    # Method that calls range_date_data passing the start and end dates of all available data
    def all_time_data(self):
        date_start_string = self.data[0].date
        date_end_string = self.data[-1].date
        self.range_date_data(date_start_string, date_end_string, True)